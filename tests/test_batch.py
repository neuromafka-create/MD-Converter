from pathlib import Path

from docx import Document
from openpyxl import Workbook

from app.core.batch import BatchConverter
from app.core.scanner import scan_paths
from app.models import ConvertOptions, FileJob


def test_scanner_and_batch(tmp_path: Path):
    root = tmp_path / "docs"
    sub = root / "sub"
    sub.mkdir(parents=True)

    d = Document()
    d.add_paragraph("Hello")
    d.save(root / "a.docx")

    wb = Workbook()
    wb.active.append(["x", "y"])
    wb.active.append([1, 2])
    wb.save(sub / "b.xlsx")

    # junk file should be ignored
    (root / "readme.txt").write_text("nope", encoding="utf-8")
    # office lock file
    Document().save(root / "~$lock.docx")

    jobs = scan_paths([root], recursive=True)
    names = {j.source.name for j in jobs}
    assert "a.docx" in names
    assert "b.xlsx" in names
    assert "~$lock.docx" not in names
    assert "readme.txt" not in names

    out = tmp_path / "md"
    options = ConvertOptions(preserve_structure=True)
    summary = BatchConverter(jobs, out, options).run_sync()

    assert summary.succeeded == 2
    assert summary.failed == 0
    assert (out / "a.md").exists()
    assert (out / "sub" / "b.md").exists()


def test_batch_isolates_errors(tmp_path: Path):
    good = tmp_path / "good.docx"
    Document().save(good)
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"not a real docx")

    jobs = [
        FileJob(source=good, relative=Path("good.docx")),
        FileJob(source=bad, relative=Path("bad.docx")),
    ]
    summary = BatchConverter(jobs, tmp_path / "out", ConvertOptions()).run_sync()
    assert summary.succeeded == 1
    assert summary.failed == 1
