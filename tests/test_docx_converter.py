from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.converters.docx_converter import DocxConverter
from app.models import ConvertOptions, FileJob


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("Intro paragraph with plain text.")
    doc.add_heading("Section", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Bold and ")
    run.bold = True
    run2 = p.add_run("italic")
    run2.italic = True

    doc.add_paragraph("Item A", style="List Bullet")
    doc.add_paragraph("Item B", style="List Bullet")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    doc.save(path)


def test_docx_conversion(tmp_path: Path):
    src = tmp_path / "sample.docx"
    _make_docx(src)
    out_dir = tmp_path / "out"

    job = FileJob(source=src, relative=Path("sample.docx"))
    options = ConvertOptions(yaml_frontmatter=True, collapse_blank_lines=True)
    result = DocxConverter().convert(job, out_dir, options)

    assert result.ok, result.error
    assert result.outputs
    md = result.outputs[0].read_text(encoding="utf-8")
    assert "---" in md
    assert "# sample" in md or "# Main Title" in md or "Main Title" in md
    assert "Intro paragraph" in md
    assert "| Name |" in md or "Name" in md
    assert "Alpha" in md
    assert result.bytes_out > 0
    assert result.bytes_out < result.bytes_in
